#!/usr/bin/env python3

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT_DOCX = ROOT / "manuscript" / "clinical_translational_oncology_submission_with_figures_v1.docx"
OUTPUT_DOCX = ROOT / "manuscript" / "clinical_translational_oncology_submission_with_figures_v2.docx"
TABLE1_CSV = ROOT / "results/tables/table1_overall_clinical_characteristics.csv"


DATA_AVAILABILITY = (
    "The Olink NPX matrix and de-identified clinical metadata are available from the "
    "corresponding author upon reasonable request, subject to institutional and patient-privacy "
    "requirements. Public TCGA-STAD data were downloaded from the UCSC Xena GDC hub on "
    "19 August 2026: STAR TPM expression matrix, "
    "https://gdc.xenahubs.net/download/TCGA-STAD.star_tpm.tsv.gz; clinical annotations, "
    "https://gdc.xenahubs.net/download/TCGA-STAD.clinical.tsv.gz; and survival annotations, "
    "https://gdc.xenahubs.net/download/TCGA-STAD.survival.tsv.gz."
)

CODE_AVAILABILITY = (
    "The analysis code is provided as Supplementary Code 1, including scripts for primary "
    "Olink analysis, clinical robustness analysis, TCGA-STAD support analysis and figure export. "
    "The code excludes raw participant-level NPX and clinical data, which are subject to "
    "institutional and patient-privacy restrictions."
)

AUTHOR_CONTRIBUTIONS = (
    "S.D. and P.C. conceived and designed the study. S.D. performed data curation, "
    "bioinformatics analysis, visualization and manuscript drafting. X.D. contributed to "
    "clinical data collection and sample organization. Q.C. contributed to sample processing "
    "and data verification. P.C. supervised the study, interpreted the results and critically "
    "revised the manuscript. All authors read and approved the final manuscript."
)

STATISTICAL_ANALYSIS = (
    "All analyses were performed in R version 4.4.1. The main packages included "
    "readxl 1.4.5, dplyr 1.1.4, tidyr 1.3.1, limma 3.60.3, glmnet 4.1-10, "
    "pROC 1.19.0.1, clusterProfiler 4.12.0, org.Hs.eg.db 3.19.1, ggplot2 4.0.3, "
    "ggrepel 0.9.5, patchwork 1.2.0, svglite 2.2.2, ragg 1.3.2, data.table 1.15.4, "
    "curl 5.2.1 and survival 3.7-0. Two-sided tests were used unless otherwise "
    "specified. FDR < 0.05 was considered statistically significant for high-dimensional "
    "analyses."
)

NAME_MAP = {
    "CA199": "CA19-9",
    "CA724": "CA72-4",
    "eGFR": "eGFR",
    "WBC": "WBC",
    "RBC": "RBC",
}


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def set_cell_text(cell, text, bold=False, size=6.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)


def clean_summary(text):
    text = str(text)
    text = text.replace("乡村", "Rural").replace("城镇", "Urban")
    text = text.replace("NA", "Not available")
    text = re.sub(r";\\s*", "; ", text)
    return text


def variable_label(name):
    return NAME_MAP.get(name, name.replace("_", " "))


def insert_after(paragraph, element):
    paragraph._p.addnext(element)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def build_table1(doc):
    with TABLE1_CSV.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))

    headers = ["Variable", "Overall", "Missing", "Healthy", "T1-T2", "T3-T4", "P value"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    widths = [0.78, 1.05, 0.43, 1.12, 1.12, 1.12, 0.48]
    for idx, header in enumerate(headers):
      cell = table.rows[0].cells[idx]
      set_cell_text(cell, header, bold=True, size=6.5)
      cell.width = Inches(widths[idx])
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        values = [
            variable_label(row["Variable"]),
            clean_summary(row["Overall"]),
            row["Missing"],
            clean_summary(row["Group_H"]),
            clean_summary(row["Group_E"]),
            clean_summary(row["Group_L"]),
            row["P_value_display"],
        ]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, size=6.1)
            cells[idx].width = Inches(widths[idx])
            if idx == 0:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif idx in (2, 6):
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    return table


def main():
    doc = Document(INPUT_DOCX)

    replacements = {
        "[Shaorui Ding1, Xiang Dong2, Qingning Chang1, Ping Chen1]": (
            "Shaorui Ding1, Xiang Dong2, Qingning Chang1, Ping Chen1"
        ),
        "The Olink NPX matrix and de-identified clinical metadata are available from the corresponding author upon reasonable request, subject to institutional and patient-privacy requirements. TCGA-STAD data were obtained from UCSC Xena/GDC [insert final URLs and download dates].": DATA_AVAILABILITY,
        "The analysis was implemented in R scripts for primary Olink analysis, clinical robustness analysis, TCGA-STAD support analysis and figure export. [Insert repository link or upon-request code-availability statement.]": CODE_AVAILABILITY,
        "[Use author initials. Example: X.X. and Y.Y. designed the study; X.X. collected samples; Y.Y. performed analysis; X.X. and Y.Y. drafted the manuscript; all authors reviewed and approved the final manuscript.]": AUTHOR_CONTRIBUTIONS,
        "All analyses were performed in R [insert version]. The main packages included readxl, dplyr, tidyr, limma, glmnet, pROC, clusterProfiler, org.Hs.eg.db, ggplot2, ggrepel, patchwork, svglite and ragg [insert package versions]. Two-sided tests were used unless otherwise specified. FDR < 0.05 was considered statistically significant for high-dimensional analyses.": STATISTICAL_ANALYSIS,
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replacements:
            paragraph.text = replacements[text]
        elif "[Insert IRB name, approval number, consent procedure and sample collection protocol.]" in text:
            paragraph.text = text.replace(
                "[Insert IRB name, approval number, consent procedure and sample collection protocol.]",
                "The study was approved by the Medical Research Ethics Review Committee of the General Hospital of Ningxia Medical University (approval number KYLL-2026-0065), and written informed consent was obtained from all participants or their legal guardians.",
            )

    # Remove the artificial-intelligence declaration section entirely.
    for idx in range(len(doc.paragraphs) - 1, -1, -1):
        if doc.paragraphs[idx].text.strip() == "Use of artificial intelligence tools":
            if idx + 1 < len(doc.paragraphs):
                delete_paragraph(doc.paragraphs[idx + 1])
            delete_paragraph(doc.paragraphs[idx])
            break

    target_para = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Continuous variables are summarized as median"):
            target_para = paragraph
            break
    if target_para is None:
        raise RuntimeError("Could not locate Table 1 legend paragraph.")

    table = build_table1(doc)
    insert_after(target_para, table._tbl)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
