#!/usr/bin/env python3

from pathlib import Path
import subprocess

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript" / "clinical_translational_oncology_submission_draft_v1.docx"
OUT = ROOT / "manuscript" / "clinical_translational_oncology_submission_with_figures_v1.docx"
FIG_TIFF = ROOT / "results" / "submission_figures" / "composite_tiff_600dpi"
EMBED_DIR = ROOT / "manuscript" / "embedded_figure_png"


FIGURES = [
    (
        "Fig. 1 Cohort design and global Olink Reveal profile",
        FIG_TIFF / "Fig1_cohort_qc_global_profile.tiff",
        6.35,
    ),
    (
        "Fig. 2 Differential and ordered plasma protein changes",
        FIG_TIFF / "Fig2_differential_protein_landscape.tiff",
        6.35,
    ),
    (
        "Fig. 3 Compact protein signatures for T1-T2 detection and invasion-depth stratification",
        FIG_TIFF / "Fig3_stage_progression_signature.tiff",
        6.35,
    ),
    (
        "Fig. 4 Clinical robustness, calibration and patient-only marker context",
        FIG_TIFF / "Fig4_sparse_diagnostic_models.tiff",
        6.35,
    ),
    (
        "Fig. 5 Biological-process context of candidate plasma protein sets",
        FIG_TIFF / "Fig5_go_biological_context.tiff",
        6.35,
    ),
]


def convert_to_png(tiff_path: Path) -> Path:
    EMBED_DIR.mkdir(parents=True, exist_ok=True)
    png_path = EMBED_DIR / (tiff_path.stem + ".png")
    subprocess.run(
        ["sips", "-s", "format", "png", str(tiff_path), "--out", str(png_path)],
        check=True,
        stdout=subprocess.DEVNULL,
    )
    return png_path


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_figure_after(paragraph, png_path: Path, width_in: float):
    spacer = insert_paragraph_after(paragraph)
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(3)
    run = spacer.add_run()
    run.add_picture(str(png_path), width=Inches(width_in))
    return spacer


def main():
    if not MANUSCRIPT.exists():
        raise FileNotFoundError(MANUSCRIPT)
    for _, fig, _ in FIGURES:
        if not fig.exists():
            raise FileNotFoundError(fig)

    doc = Document(str(MANUSCRIPT))

    title_to_png = {title: convert_to_png(path) for title, path, _ in FIGURES}
    title_to_width = {title: width for title, _, width in FIGURES}

    paragraphs = list(doc.paragraphs)
    inserted = []
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text not in title_to_png:
            continue

        # Insert after the legend body paragraph immediately following the figure heading.
        anchor = paragraph
        if idx + 1 < len(paragraphs) and paragraphs[idx + 1].text.strip():
            anchor = paragraphs[idx + 1]

        fig_para = insert_figure_after(anchor, title_to_png[text], title_to_width[text])
        fig_para.paragraph_format.keep_with_next = False
        inserted.append(text)

        # Page-break after each figure except the last; keeps captions and figures readable.
        if text != FIGURES[-1][0]:
            pb = insert_paragraph_after(fig_para)
            pb.add_run().add_break(WD_BREAK.PAGE)

    missing = [title for title, _, _ in FIGURES if title not in inserted]
    if missing:
        raise RuntimeError("Figure anchors not found: " + "; ".join(missing))

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
